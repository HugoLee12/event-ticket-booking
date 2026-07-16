"""Sinh docx thân báo cáo (TOC tự động + 9 mục) từ 9 file Markdown nguồn.

Chạy: pip install python-docx  &&  python build_report.py
Kết quả: report-body.docx (chỉ gồm mục lục tự động + thân 9 mục, KHÔNG bìa - tự thêm sau).

Markdown là nguồn chân lý; docx sinh tự động. Đổi số liệu/câu chữ chỉ cần sửa .md rồi chạy lại.
Chạy trực tiếp file này cũng chạy luôn self-check (đọc lại docx vừa sinh, assert các bất biến).
"""

import os
import re
import sys
import glob
import hashlib
import urllib.request

if hasattr(sys.stdout, "reconfigure"):  # console Windows mặc định cp1252, ép UTF-8 để in tiếng Việt
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "report-body.docx")

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
CODE_FILL = "F2F2F2"
PLACEHOLDER_PREFIX = "[Sơ đồ: chưa render được"  # dùng cho cả chèn lẫn self-check
KROKI_URL = "https://kroki.io/mermaid/png"


# ---------- helpers XML cấp thấp python-docx không expose ----------

def _set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _add_field(run, instr, placeholder_text):
    """Chèn một Word field (TOC, PAGE...) vào run: begin -> instrText -> separate -> text -> end."""
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = placeholder_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, txt, end):
        r.append(el)


# ---------- thiết lập style & khung tài liệu ----------

def _setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(13)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)

    for name, size, pb in (("Heading 1", 16, True), ("Heading 2", 14, False)):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.page_break_before = pb
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def _add_footer_page_number(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(p.add_run(), "PAGE", "1")


def _add_toc(doc):
    h = doc.add_paragraph("Mục lục")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(16)
    note = doc.add_paragraph("(Bấm Ctrl+A rồi F9 để cập nhật mục lục và số trang.)")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(11)
    _add_field(doc.add_paragraph().add_run(), 'TOC \\o "1-2" \\h \\z \\u', "Mục lục sẽ hiện ở đây sau khi bấm F9.")


# ---------- parser markdown (tập con 9 file thực dùng) ----------

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_inline_runs(paragraph, text):
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(11)
        else:
            paragraph.add_run(piece)


def _content_width(doc):
    sec = doc.sections[0]
    return sec.page_width - sec.left_margin - sec.right_margin


def _max_height(doc):
    sec = doc.sections[0]
    return sec.page_height - sec.top_margin - sec.bottom_margin


def _insert_picture(doc, abspath):
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    shape = run.add_picture(abspath, width=_content_width(doc))
    max_h = _max_height(doc)
    if shape.height > max_h:  # sơ đồ quá cao -> co theo chiều cao, giữ tỉ lệ
        shape.width = int(shape.width * max_h / shape.height)
        shape.height = max_h


def _kroki_png(mermaid_text):
    req = urllib.request.Request(
        KROKI_URL,
        data=mermaid_text.encode("utf-8"),
        headers={"Content-Type": "text/plain", "User-Agent": "Mozilla/5.0"},
    )
    with urllib.request.urlopen(req, timeout=20) as resp:
        return resp.read()


def _render_mermaid(doc, block_text, warnings):
    """Khối mermaid là nguồn duy nhất: hash nội dung -> cache img/mermaid-<hash8>.png.
    Cache hit -> chèn inline; miss -> render qua Kroki rồi cache; lỗi mạng -> placeholder + warning."""
    digest = hashlib.sha1(block_text.encode("utf-8")).hexdigest()[:8]
    rel = "img/mermaid-%s.png" % digest
    abspath = os.path.join(HERE, "img", "mermaid-%s.png" % digest)
    if not os.path.exists(abspath):
        try:
            data = _kroki_png(block_text)
            os.makedirs(os.path.dirname(abspath), exist_ok=True)
            with open(abspath, "wb") as fh:
                fh.write(data)
        except Exception as exc:  # mất mạng và cache chưa có -> không fail build
            warnings.append((rel, str(exc)))
            p = doc.add_paragraph("%s %s]" % (PLACEHOLDER_PREFIX, rel))
            p.runs[0].italic = True
            return False
    _insert_picture(doc, abspath)
    return True


def _add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, CODE_FILL)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)  # bỏ đoạn rỗng mặc định
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = CODE_FONT
        r.font.size = Pt(10)


def _add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].paragraphs[0].text = ""
            _add_inline_runs(cells[i].paragraphs[0], row[i] if i < len(row) else "")
            for rn in cells[i].paragraphs[0].runs:
                rn.font.size = Pt(11)


def _parse_table(md_lines, start):
    """Đọc một bảng markdown từ start; trả (rows, next_index). rows[0] là header."""
    def split_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    rows = [split_row(md_lines[start])]
    i = start + 1  # dòng phân cách |---|---|
    i += 1
    while i < len(md_lines) and md_lines[i].lstrip().startswith("|"):
        rows.append(split_row(md_lines[i]))
        i += 1
    return rows, i


def _render_markdown(doc, md_lines, warnings):
    n_tables = 0
    n_mermaid = 0
    n_placeholder = 0
    i = 0
    while i < len(md_lines):
        line = md_lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):  # code fence
            lang = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < len(md_lines) and not md_lines[i].strip().startswith("```"):
                block.append(md_lines[i])
                i += 1
            i += 1  # bỏ dòng đóng ```
            if lang == "mermaid":  # khối mermaid = nguồn duy nhất, render inline tại chỗ
                n_mermaid += 1
                if not _render_mermaid(doc, "\n".join(block), warnings):
                    n_placeholder += 1
            else:
                _add_code_block(doc, block)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(md_lines) and set(md_lines[i + 1].strip()) <= set("|-: "):
            rows, i = _parse_table(md_lines, i)
            _add_table(doc, rows)
            n_tables += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    return n_tables, n_mermaid, n_placeholder


# ---------- build ----------

def build():
    doc = Document()
    _setup_styles(doc)
    _add_footer_page_number(doc)
    _add_toc(doc)

    files = sorted(glob.glob(os.path.join(HERE, "0[0-9]-*.md")))
    total_tables = 0
    total_mermaid = 0
    total_placeholder = 0
    warnings = []
    for f in files:
        with open(f, encoding="utf-8") as fh:
            md_lines = fh.read().splitlines()
        t, m, ph = _render_markdown(doc, md_lines, warnings)
        total_tables += t
        total_mermaid += m
        total_placeholder += ph

    doc.save(OUT_PATH)
    print("Đã sinh %s (%d file md, %d bảng, %d sơ đồ mermaid)." % (
        os.path.basename(OUT_PATH), len(files), total_tables, total_mermaid))
    if warnings:
        print("CẢNH BÁO: %d sơ đồ chưa render được (mất mạng và cache chưa có), đã chèn placeholder:" % len(warnings))
        for rel, err in warnings:
            print("  - %s (%s)" % (rel, err))
    return total_tables, total_mermaid


# ---------- self-check: đọc lại docx, assert bất biến ở đầu ra ----------

def self_check(src_tables, src_mermaid):
    doc = Document(OUT_PATH)
    h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]

    # (a) đúng 9 Heading 1 mang nhãn 1..9
    assert len(h1) == 9, "Cần đúng 9 Heading 1, thấy %d" % len(h1)
    for idx, text in enumerate(h1, start=1):
        assert text.startswith("%d." % idx), "Heading 1 thứ %d sai nhãn: %r" % (idx, text)

    # (b) đủ nhiều Heading 2
    assert len(h2) >= 30, "Quá ít Heading 2: %d" % len(h2)

    # (c) số bảng Word >= bảng md nguồn; bảng STRIDE có hàng tiêu đề đậm
    assert len(doc.tables) >= src_tables, "Thiếu bảng: %d < %d" % (len(doc.tables), src_tables)
    stride = None
    for t in doc.tables:
        head = " ".join(c.text for c in t.rows[0].cells)
        if "STRIDE" in head:
            stride = t
            break
    assert stride is not None, "Không tìm thấy bảng STRIDE"
    assert all(any(r.bold for r in c.paragraphs[0].runs) for c in stride.rows[0].cells), \
        "Hàng tiêu đề bảng STRIDE không đậm"

    # (d) mỗi khối mermaid nguồn -> đúng 1 ảnh inline HOẶC 1 placeholder (không mất)
    n_pics = len(doc.inline_shapes)
    n_placeholder = sum(1 for p in doc.paragraphs if p.text.startswith(PLACEHOLDER_PREFIX))
    assert n_pics + n_placeholder >= src_mermaid, \
        "Mất sơ đồ mermaid: %d ảnh + %d placeholder < %d" % (n_pics, n_placeholder, src_mermaid)

    # (e) một câu mẫu từ mục 1 xuất hiện nguyên văn (giữ dấu tiếng Việt)
    sample = "một ghế chỉ được bán cho tối đa một người, trong mọi tình huống đồng thời"
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert sample in all_text, "Không thấy câu mẫu mục 1 trong docx (nội dung/dấu bị mất?)"

    # (f) có đúng 1 field TOC
    xml = doc.element.xml
    assert xml.count("TOC ") >= 1, "Không thấy field TOC"

    print("Self-check OK: 9 H1, %d H2, %d bảng (STRIDE header đậm), %d ảnh + %d placeholder, TOC + câu mẫu tiếng Việt." % (
        len(h2), len(doc.tables), n_pics, n_placeholder))


if __name__ == "__main__":
    tables, mermaid = build()
    self_check(tables, mermaid)
